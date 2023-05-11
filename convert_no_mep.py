import os
import openpyxl
from docx import Document
from docx.oxml.ns import qn
import glob
import re 
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import time 

fichier_excel = "items.xlsx"
dossier_modeles = "./input_files"
dossier_sortie = "output"

def paragraph_replace_text(paragraph, research, replace_str):
    remplacer_texte_lien_hypertexte(paragraph, research, replace_str)
    replace_text(paragraph, research, replace_str)

def replace_text(paragraph, research, replace_str):

    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        regex = re.compile(str(research), re.IGNORECASE)
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break
    
        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph


def remplacer_texte_lien_hypertexte(paragraphe, recherche, remplacement):
    recherche_str = str(recherche)
    regex = re.compile(r'(?<![a-zA-Z0-9\-\.@])' + re.escape(recherche_str) + r'(?![a-zA-Z0-9\-\.@])', re.IGNORECASE)
    ns = {"w": paragraphe._element.nsmap["w"]}

    for hyperlink in paragraphe._element.iter(f"{{{ns['w']}}}hyperlink"):
        runs = hyperlink.findall(f"{{{ns['w']}}}r")
        merged_text = "".join(run.text for run in runs)
        merged_text = regex.sub(str(remplacement), merged_text)

        # Séparer le texte fusionné en plusieurs éléments de texte
        new_runs = []
        start_index = 0
        for match in regex.finditer(merged_text):
            match_start = match.start()
            match_end = match.end()
            if match_start > start_index:
                text_before = merged_text[start_index:match_start]
                new_run = parse_xml(f"<w:r {nsdecls('w')}><w:t {nsdecls('w')}>{text_before}</w:t></w:r>")
                new_runs.append(new_run)
            replaced_text = match.group()
            new_run = parse_xml(f"<w:r {nsdecls('w')}><w:t {nsdecls('w')} xml:space='preserve'>{replaced_text}</w:t></w:r>")
            new_run.run_properties = runs[0].run_properties  # Copy the run properties from the first run
            new_runs.append(new_run)
            start_index = match_end
        if start_index < len(merged_text):
            text_after = merged_text[start_index:]
            new_run = parse_xml(f"<w:r {nsdecls('w')}><w:t {nsdecls('w')}>{text_after} </w:t></w:r>")
            new_runs.append(new_run)

        # Remplacer les runs existants par les nouveaux runs
        for run in runs:
            hyperlink.remove(run)
        for new_run in new_runs:
            hyperlink.append(new_run)



# Fonction pour remplacer le texte dans les en-têtes et pieds de page
def remplacer_texte_entete_pied_page(entete_pied_page, recherche, remplacement):
    
    for paragraph in entete_pied_page.paragraphs:
        paragraph_replace_text(paragraph, recherche, remplacement)
    for paragraphe in entete_pied_page.tables:
        for row in paragraphe.rows:
            for cell in row.cells:
                for paragraphe in cell.paragraphs:
                    paragraph_replace_text(paragraphe, recherche, remplacement)


# Fonction pour créer un document Word pour chaque utilisateur
def creer_documents_utilisateurs(fichier_excel, repertoire_modeles, dossier_sortie):
    # Chargez le fichier Excel contenant les valeurs à rechercher et à remplacer
    wb = openpyxl.load_workbook(fichier_excel)
    ws = wb.active
    # Récupérez les valeurs de recherche et les valeurs de remplacement pour chaque utilisateur
    valeurs_recherche = [cell.value for cell in ws[1] if cell.column > 1]
    valeurs_utilisateurs = {}
    
    for row in ws.iter_rows(min_row=2, values_only=True, min_col=2):
        recherche = row[0]
        for i, remplacement in enumerate(row[1:], 1):
            mots_recherche = str(recherche).split(";")  # Sépare les mots à rechercher par point-virgule
            for mot_recherche in mots_recherche:
                if i not in valeurs_utilisateurs:
                    valeurs_utilisateurs[i] = {}
                valeurs_utilisateurs[i][mot_recherche] = remplacement

    # Créez un document Word pour chaque utilisateur
    for i, valeurs_remplacement in valeurs_utilisateurs.items():
        # Créez un dossier unique pour chaque utilisateur
        premiere_valeur = list(valeurs_remplacement.values())[0]
        dossier_utilisateur = os.path.join(dossier_sortie, str(premiere_valeur))
        os.makedirs(dossier_utilisateur, exist_ok=True)

        # Ouvrez le modèle Word et remplacez les valeurs
        modeles = glob.glob(os.path.join(repertoire_modeles, "*.docx"))

        for modele in modeles:
            print("Traitement du fichier: " + modele)
            doc = Document(modele)
            for paragraphe in doc.paragraphs:
                for recherche, remplacement in valeurs_remplacement.items():
                    paragraph_replace_text(paragraphe, recherche, remplacement)
    

            # Parcourez les cellules de tableau et remplacez les valeurs
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraphe in cell.paragraphs:
                            for recherche, remplacement in valeurs_remplacement.items():
                                paragraph_replace_text(paragraphe, recherche, remplacement)
            # Parcourez les en-têtes et pieds de page et remplacez les valeurs
            for section in doc.sections:
                header = section.header
                footer = section.footer
                first_header = section.first_page_header

                if first_header is not None:
                    for recherche, remplacement in valeurs_remplacement.items():
                        remplacer_texte_entete_pied_page(first_header, recherche, remplacement)
                if header is not None:
                    for recherche, remplacement in valeurs_remplacement.items():
                        remplacer_texte_entete_pied_page(header, recherche, remplacement)
                if footer is not None:
                    for recherche, remplacement in valeurs_remplacement.items():
                        remplacer_texte_entete_pied_page(footer, recherche, remplacement)

            # Sauvegardez le document rempli
            fichier_sortie = os.path.join(dossier_utilisateur, os.path.basename(modele))
            doc.save(fichier_sortie)

start_time = time.time()
creer_documents_utilisateurs(fichier_excel, dossier_modeles, dossier_sortie)
end_time = time.time()
execution_time = end_time - start_time
print(f"Temps d'exécution : {execution_time} secondes")
